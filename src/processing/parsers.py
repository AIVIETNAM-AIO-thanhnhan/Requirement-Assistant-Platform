"""
parsers.py — Bộ parser tài liệu hợp nhất cho PDF, DOCX, XLSX.

Mỗi parser nhận đường dẫn file -> trả về text thô (str).
Dùng chung cho build_chunks.py và build_catalogue.py.

  parse_pdf  : trích text PDF dạng text (KHÔNG OCR; PDF scan dùng ocr_pdf_hq.py riêng)
  parse_docx : đọc cả đoạn văn (paragraph) VÀ bảng (table) trong file Word
  parse_xlsx : đọc mọi sheet, mỗi dòng nối các ô bằng " | "

Cài: pip install pymupdf python-docx openpyxl
"""

from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from openpyxl import load_workbook


def parse_pdf(path: Path) -> str:
    """Trích text từ PDF dạng text. Trả về '' nếu là PDF scan (không có text)."""
    doc = fitz.open(str(path))
    text = "\n".join(page.get_text() for page in doc)
    doc.close()
    return text


def parse_docx(path: Path) -> str:
    """Đọc đoạn văn + bảng trong file Word. (Bản cũ bỏ sót bảng -> đã sửa.)"""
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def parse_xlsx(path: Path) -> str:
    """Đọc mọi sheet trong file Excel. Mỗi dòng nối các ô bằng ' | '."""
    wb = load_workbook(str(path), data_only=True, read_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"# Sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                parts.append(" | ".join(cells))
    wb.close()
    return "\n".join(parts)


# Bảng tra parser theo đuôi file
PARSERS = {
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".xlsx": parse_xlsx,
}

SUPPORTED = set(PARSERS) | {".txt"}


def parse(path: Path) -> str:
    """Parse 1 file theo đuôi. .txt đọc trực tiếp. Đuôi lạ -> trả về ''."""
    suffix = path.suffix.lower()
    if suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")
    fn = PARSERS.get(suffix)
    return fn(path) if fn else ""
